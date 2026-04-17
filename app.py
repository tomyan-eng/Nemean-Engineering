import streamlit as st
import google.generativeai as genai
from PIL import Image
import io
import base64
from datetime import datetime
import PyPDF2
import docx
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from streamlit.components.v1 import html

st.set_page_config(page_title="Nemean Engineering Inspector", layout="wide")
st.title("🏗️ Nemean Engineering Inspector (Cloud Persistent)")
st.markdown("Deficiency ID · Code/Tender Compliance · Work Progress Tracking")

# ---------- Firebase REST API helpers ----------
def get_firestore_url(collection, doc_id=None):
    project_id = st.secrets["FIREBASE_PROJECT_ID"]
    base = f"https://firestore.googleapis.com/v1/projects/{project_id}/databases/(default)/documents/{collection}"
    if doc_id:
        return f"{base}/{doc_id}"
    return base

def load_all_projects():
    """Fetch all projects from Firestore"""
    url = get_firestore_url("projects")
    params = {"key": st.secrets["FIREBASE_API_KEY"]}
    try:
        response = requests.get(url, headers={"Content-Type": "application/json"}, params=params)
        if response.status_code == 200:
            data = response.json()
            projects = {}
            for doc in data.get("documents", []):
                doc_id = doc["name"].split("/")[-1]
                fields = doc["fields"]
                proj_data = {
                    "name": fields.get("name", {}).get("stringValue", doc_id),
                    "visits": json.loads(fields.get("visits", {}).get("stringValue", "[]"))
                }
                projects[doc_id] = proj_data
            return projects
        else:
            st.error(f"Failed to load projects: {response.text}")
            return {}
    except Exception as e:
        st.error(f"Error loading: {e}")
        return {}

def save_project(project_name, project_data):
    """Save or update a project in Firestore"""
    doc_id = project_name.replace(" ", "_").replace("/", "_")
    url = get_firestore_url("projects", doc_id) + f"?key={st.secrets['FIREBASE_API_KEY']}"
    fields = {
        "name": {"stringValue": project_name},
        "visits": {"stringValue": json.dumps(project_data.get("visits", []))},
        "updated_at": {"stringValue": datetime.now().isoformat()}
    }
    body = {"fields": fields}
    response = requests.patch(url, headers={"Content-Type": "application/json"}, json=body)
    return response.status_code in [200, 201]

def delete_project(project_name):
    doc_id = project_name.replace(" ", "_").replace("/", "_")
    url = get_firestore_url("projects", doc_id) + f"?key={st.secrets['FIREBASE_API_KEY']}"
    response = requests.delete(url, headers={"Content-Type": "application/json"})
    return response.status_code == 200

# ---------- API Key ----------
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    key_configured = True
except:
    key_configured = False

# ---------- Session State ----------
if "images_data" not in st.session_state:
    st.session_state["images_data"] = []
if "reference_text" not in st.session_state:
    st.session_state["reference_text"] = ""
if "report_template" not in st.session_state:
    st.session_state["report_template"] = None
if "persistent_prompt" not in st.session_state:
    st.session_state["persistent_prompt"] = """You are an engineering compliance assistant.
Based on the provided reference, analyze each construction photo.
Output in this exact format:
- Deficiency: [specific observation]
- Compliance: [Compliant / Non-compliant / N/A]
- Reference: [cite the relevant clause/section from the reference]
- Remedy: [action if non-compliant]
- Severity: [Low/Medium/High/Critical]

Be concise. Do not add extra commentary."""
if "projects" not in st.session_state:
    st.session_state.projects = load_all_projects()
if "current_project" not in st.session_state:
    st.session_state.current_project = None
if "current_visit_date" not in st.session_state:
    st.session_state.current_visit_date = None

# ---------- Sidebar ----------
with st.sidebar:
    st.header("⚙️ Configuration")
    
    if not key_configured:
        api_key_input = st.text_input("Gemini API Key", type="password")
        if api_key_input:
            genai.configure(api_key=api_key_input)
            key_configured = True
    else:
        st.success("✅ Gemini API ready")
    
    st.markdown("---")
    mode = st.radio("Main mode", ["Deficiency & Compliance", "Work Progress Tracking"])
    
    st.markdown("---")
    
    if mode == "Deficiency & Compliance":
        st.subheader("📜 Reference (Codes / Tender)")
        ref_type = st.radio("Reference type", ["Building Codes (OBC/OFC)", "Tender / CCDC Documents"])
        if ref_type == "Building Codes (OBC/OFC)":
            default_ref = """
Ontario Building Code 2012: Section 9 (Housing), 5 (Environmental Separation)
Ontario Fire Code O.Reg 213/07
CSA A440.2 (Windows/Doors), CSA A23.1 (Concrete)
ASTM C920 (Sealants), ASTM D6163 (Modified bitumen)
OBC 9.26 (Roof coverings), 9.27 (Cladding), 5.4 (Moisture protection)
"""
            st.session_state["reference_text"] = st.text_area("Reference codes", value=default_ref.strip(), height=200)
        else:
            ref_files = st.file_uploader("Upload PDF/DOCX/TXT (tender, CCDC)", type=["pdf","docx","txt"], accept_multiple_files=True)
            if ref_files:
                combined = ""
                for f in ref_files:
                    try:
                        if f.type == "application/pdf":
                            reader = PyPDF2.PdfReader(f)
                            for page in reader.pages:
                                combined += page.extract_text() or ""
                        elif f.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            doc = docx.Document(f)
                            combined += "\n".join([p.text for p in doc.paragraphs])
                        else:
                            combined += f.read().decode("utf-8")
                    except Exception as e:
                        st.error(f"Error reading {f.name}: {e}")
                st.session_state["reference_text"] = combined
                with st.expander("Preview extracted text"):
                    st.text(combined[:1000] + ("..." if len(combined)>1000 else ""))
            else:
                st.info("Upload tender/CCDC documents for compliance analysis.")
    
    else:  # Progress Tracking mode
        st.subheader("📁 Project Management (Cloud)")
        project_names = list(st.session_state.projects.keys())
        selected = st.selectbox("Select or create project", ["-- New Project --"] + project_names)
        if selected == "-- New Project --":
            new_name = st.text_input("New project name")
            if new_name and st.button("Create Project"):
                if new_name not in st.session_state.projects:
                    st.session_state.projects[new_name] = {"visits": []}
                    save_project(new_name, st.session_state.projects[new_name])
                    st.session_state.current_project = new_name
                    st.rerun()
        else:
            st.session_state.current_project = selected
            st.success(f"Current: **{selected}**")
            visits = st.session_state.projects[selected].get("visits", [])
            if visits:
                visit_dates = [v["date"] for v in visits]
                st.selectbox("Previous visits (for reference)", visit_dates, key="prev_visit")
            if st.button("➕ Start new site visit"):
                st.session_state.current_visit_date = datetime.now().strftime("%Y-%m-%d %H:%M")
                st.session_state.images_data = []
                st.rerun()
            if st.session_state.images_data and st.button("💾 Save current visit to cloud"):
                visit_data = {
                    "date": st.session_state.current_visit_date or datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "photos_b64": [],
                    "notes": [],
                    "analyses": []
                }
                for img_data in st.session_state.images_data:
                    buff = io.BytesIO()
                    img_data["image"].save(buff, format="JPEG")
                    b64 = base64.b64encode(buff.getvalue()).decode()
                    visit_data["photos_b64"].append(b64)
                    visit_data["notes"].append(img_data["notes"])
                    visit_data["analyses"].append(img_data["analysis"])
                st.session_state.projects[st.session_state.current_project]["visits"].append(visit_data)
                save_project(st.session_state.current_project, st.session_state.projects[st.session_state.current_project])
                st.success("Visit saved to cloud!")
            if st.button("🗑️ Delete this project"):
                delete_project(selected)
                del st.session_state.projects[selected]
                st.session_state.current_project = None
                st.rerun()
        st.session_state["reference_text"] = st.text_area("Project scope / schedule notes (optional)", height=100)
    
    st.markdown("---")
    st.subheader("✏️ Analysis Prompt")
    prompt_editor = st.text_area("Edit prompt (persistent)", value=st.session_state["persistent_prompt"], height=200)
    if prompt_editor != st.session_state["persistent_prompt"]:
        st.session_state["persistent_prompt"] = prompt_editor
    
    st.markdown("---")
    st.subheader("📝 Report Template")
    template_file = st.file_uploader("Upload HTML or DOCX template", type=["html","docx"])
    if template_file:
        try:
            if template_file.name.endswith(".html"):
                template_html = template_file.read().decode("utf-8")
                st.session_state["report_template"] = template_html
            else:
                doc = docx.Document(template_file)
                template_html = "\n".join([p.text for p in doc.paragraphs])
                st.session_state["report_template"] = template_html
            st.success("Template loaded. Use {{DATE}}, {{PHOTO_1}}, {{ANALYSIS_1}}, etc.")
        except Exception as e:
            st.error(f"Template error: {e}")
    else:
        st.info("No template = default HTML report.")
    
    if st.button("🔄 Refresh projects from cloud"):
        st.session_state.projects = load_all_projects()
        st.rerun()
    
    if st.button("🗑️ Clear current photos"):
        st.session_state.images_data = []
        st.rerun()

# ---------- Helper Functions ----------
def analyze_single_image(image_pil, user_notes, reference_text, prompt_text):
    if not key_configured:
        return "❌ API key missing."
    try:
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        full_prompt = f"{prompt_text}\n\nReference:\n{reference_text}\n\nInspector notes: {user_notes}\nAnalyze the photo."
        response = model.generate_content([full_prompt, image_pil])
        return response.text
    except Exception as e:
        return f"⚠️ Error: {e}"

def batch_analyze_all(images_list, reference_text, prompt_text, progress_bar):
    total = len(images_list)
    results = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        future_to_idx = {}
        for idx, img_data in enumerate(images_list):
            future = executor.submit(analyze_single_image, img_data["image"], img_data["notes"], reference_text, prompt_text)
            future_to_idx[future] = idx
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                result = future.result()
            except Exception as e:
                result = f"Error: {e}"
            images_list[idx]["analysis"] = result
            results.append((idx, result))
            progress_bar.progress(len(results)/total)
    return images_list

# ---------- Main Area: Photo Upload & Analysis ----------
col1, col2 = st.columns([1,1])

with col1:
    st.subheader("📸 Upload Site Photos")
    uploaded = st.file_uploader("Select images (JPG, PNG)", type=["jpg","jpeg","png"], accept_multiple_files=True)
    if uploaded:
        for file in uploaded:
            if not any(d["file_name"] == file.name for d in st.session_state.images_data):
                img = Image.open(file).convert("RGB")
                st.session_state.images_data.append({
                    "image": img,
                    "notes": "",
                    "analysis": "",
                    "file_name": file.name
                })
        st.success(f"{len(uploaded)} photo(s) added.")
    
    if st.session_state.images_data:
        for idx, data in enumerate(st.session_state.images_data):
            with st.container():
                a,b = st.columns([1,2])
                with a:
                    st.image(data["image"], width=150, caption=data["file_name"])
                with b:
                    new_notes = st.text_area(f"Notes for {data['file_name']}", value=data["notes"], key=f"notes_{idx}")
                    data["notes"] = new_notes
                    if st.button(f"🔍 Analyze single", key=f"single_{idx}"):
                        with st.spinner("Analyzing..."):
                            data["analysis"] = analyze_single_image(
                                data["image"], data["notes"], st.session_state["reference_text"], st.session_state["persistent_prompt"]
                            )
                            st.rerun()
                    if data["analysis"]:
                        st.markdown("**Analysis:**")
                        st.markdown(data["analysis"])
                st.markdown("---")
        
        if st.button("⚡ Analyze all photos (batch)", type="primary"):
            if not st.session_state["reference_text"].strip() and mode != "Work Progress Tracking":
                st.error("Please provide reference text in sidebar.")
            else:
                progress = st.progress(0)
                with st.spinner("Batch analyzing all photos (parallel)..."):
                    updated = batch_analyze_all(
                        st.session_state.images_data,
                        st.session_state["reference_text"],
                        st.session_state["persistent_prompt"],
                        progress
                    )
                    st.session_state.images_data = updated
                st.success("All photos analyzed!")
                st.rerun()

# ---------- Report Generation ----------
with col2:
    st.subheader("📄 Generate Report")
    analyzed_count = sum(1 for d in st.session_state.images_data if d["analysis"])
    st.info(f"📊 {analyzed_count}/{len(st.session_state.images_data)} photos analyzed")
    
    report_title = st.text_input("Report title", value=f"Inspection Report – {datetime.now().strftime('%Y-%m-%d')}")
    project_info = st.text_area("Project details (address, client, inspector name)", height=80)
    
    if st.button("📝 Generate Report", type="primary") and st.session_state.images_data:
        if not key_configured:
            st.error("API key missing")
        else:
            placeholders = {
                "DATE": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "TITLE": report_title,
                "PROJECT_INFO": project_info.replace("\n", "<br>"),
                "TOTAL_PHOTOS": str(len(st.session_state.images_data))
            }
            for i, data in enumerate(st.session_state.images_data, start=1):
                buff = io.BytesIO()
                data["image"].save(buff, format="JPEG")
                img_b64 = base64.b64encode(buff.getvalue()).decode()
                placeholders[f"PHOTO_{i}"] = f'<img src="data:image/jpeg;base64,{img_b64}" style="max-width:100%">'
                placeholders[f"NOTES_{i}"] = data["notes"] or "—"
                placeholders[f"ANALYSIS_{i}"] = data["analysis"] or "Not analyzed"
                placeholders[f"FILENAME_{i}"] = data["file_name"]
            
            if st.session_state["report_template"]:
                html_output = st.session_state["report_template"]
                for key, value in placeholders.items():
                    html_output = html_output.replace(f"{{{{{key}}}}}", str(value))
            else:
                html_output = f"""
                <html><head><meta charset="UTF-8"><title>{report_title}</title>
                <style>body{{font-family:Arial;margin:40px}} .photo{{margin-bottom:30px;border-bottom:1px solid #ccc}}</style>
                </head><body>
                <h1>{report_title}</h1>
                <p><strong>Generated:</strong> {placeholders['DATE']}</p>
                <p>{placeholders['PROJECT_INFO']}</p>
                <hr>
                """
                for i, data in enumerate(st.session_state.images_data, start=1):
                    buff = io.BytesIO()
                    data["image"].save(buff, format="JPEG")
                    b64 = base64.b64encode(buff.getvalue()).decode()
                    html_output += f"""
                    <div class="photo">
                        <h3>Photo {i}: {data['file_name']}</h3>
                        <img src="data:image/jpeg;base64,{b64}" style="max-width:100%">
                        <p><strong>Notes:</strong> {data['notes'] or '—'}</p>
                        <p><strong>Analysis:</strong><br>{data['analysis'].replace(chr(10),'<br>') if data['analysis'] else 'Not analyzed'}</p>
                    </div><hr>
                    """
                html_output += "<p><i>Generated by Nemean Engineering Inspector (Cloud)</i></p></body></html>"
            
            st.session_state["final_report"] = html_output
            st.success("Report generated")
    
    if "final_report" in st.session_state:
        st.components.v1.html(st.session_state["final_report"], height=500, scrolling=True)
        b64 = base64.b64encode(st.session_state["final_report"].encode()).decode()
        st.markdown(f'<a href="data:text/html;base64,{b64}" download="{report_title.replace(" ","_")}.html">⬇️ Download HTML (print to PDF)</a>', unsafe_allow_html=True)

# ---------- Progress Mode Extra: Timeline Report ----------
if mode == "Work Progress Tracking" and st.session_state.current_project and st.session_state.projects.get(st.session_state.current_project, {}).get("visits", []):
    st.markdown("---")
    st.subheader("📈 Project Progress Timeline")
    project = st.session_state.projects[st.session_state.current_project]
    visits = project["visits"]
    if len(visits) >= 1:
        if st.button("Generate Progress Summary Report"):
            html_progress = f"<h1>Progress Report: {st.session_state.current_project}</h1>"
            html_progress += f"<p>Number of site visits: {len(visits)}</p>"
            for idx, visit in enumerate(visits):
                html_progress += f"<h3>Visit {idx+1}: {visit['date']}</h3>"
                for i, (b64, note, analysis) in enumerate(zip(visit["photos_b64"], visit["notes"], visit["analyses"])):
                    html_progress += f"""
                    <div style="margin-bottom:20px">
                        <img src="data:image/jpeg;base64,{b64}" width="300">
                        <p><strong>Notes:</strong> {note}</p>
                        <p><strong>Analysis:</strong> {analysis}</p>
                    </div>
                    """
                html_progress += "<hr>"
            st.components.v1.html(html_progress, height=600, scrolling=True)
            b64_progress = base64.b64encode(html_progress.encode()).decode()
            st.markdown(f'<a href="data:text/html;base64,{b64_progress}" download="progress_{st.session_state.current_project}.html">⬇️ Download Progress Report</a>', unsafe_allow_html=True)
    else:
        st.info("No visits saved yet. Start a new visit and save it.")

st.caption("✅ Data permanently stored in Firebase Cloud. Use 'Refresh projects from cloud' to sync across devices.")
